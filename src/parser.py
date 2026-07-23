from __future__ import annotations

import re
from pathlib import Path
from typing import BinaryIO

import pdfplumber
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def clean_text(text: str) -> str:
    """
    Remove unnecessary whitespace while preserving readable text.
    """
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def extract_text_from_pdf(file: BinaryIO) -> str:
    """
    Extract text from a PDF file.
    """
    pages: list[str] = []

    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)

    return clean_text("\n".join(pages))


def extract_text_from_docx(file: BinaryIO) -> str:
    """
    Extract paragraphs and table content from a DOCX file.
    """
    document = Document(file)
    content: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )

            if row_text:
                content.append(row_text)

    return clean_text("\n".join(content))


def extract_text_from_txt(file: BinaryIO) -> str:
    """
    Extract text from a TXT file.
    """
    raw_content = file.read()

    if isinstance(raw_content, bytes):
        text = raw_content.decode("utf-8", errors="ignore")
    else:
        text = str(raw_content)

    return clean_text(text)


def extract_text(file: BinaryIO, filename: str) -> str:
    """
    Select the correct parser based on the file extension.
    """
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            "Only PDF, DOCX and TXT files are supported."
        )

    file.seek(0)

    if extension == ".pdf":
        text = extract_text_from_pdf(file)

    elif extension == ".docx":
        text = extract_text_from_docx(file)

    else:
        text = extract_text_from_txt(file)

    if not text:
        raise ValueError(
            f"No readable text was found in {filename}."
        )

    return text